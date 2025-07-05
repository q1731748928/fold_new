from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
import os
import time

document = Document()
# docx 设置，如：字体
document.styles['Normal'].font.name = u'Calibri'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Calibri')
document.styles['Normal'].font.size = Pt(11)
document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)


paragraph = document.add_paragraph()
run = paragraph.add_run()
#run.font.name = 'Times New Roman'
run.font.name=u'SimSun'
run.font.color.rgb = RGBColor(0,0,0)
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'SimSun')

# 打开txt文件，并写入document中
NUM_MIN = 0
NUM_MAX = 50000
path = 'D:\\pythonProject\\turing_exec\\日常兼职\\小程序\\题库2\\题库2\\图片文本'
# 获取文件夹中所有的 .txt 文件
txt_files = [f for f in os.listdir(path) if f.endswith('.txt')]
# 遍历每个 .txt 文件
for txt_file in txt_files:
    # 构建完整的文件路径
    txt_file_path = os.path.join(path, txt_file)
    # r = os.path.join(path, 'Screenshot_20250424_171057_com.tencent.mm.txt')
    s = os.path.join(path, str(NUM_MAX)+"_图片题库2.docx")
    with open(txt_file_path, 'r', encoding='utf-8') as rf:
        num,last_time,add = 0, 0, 0
        for line in rf:
            num += 1
            line = line.strip()
            if NUM_MIN < num <= NUM_MAX:
                add += 1
                document.add_paragraph(line)
            elif num > NUM_MAX:
                break
# 保存文档
document.save(s)

print(num, add, '任务完成！！')
