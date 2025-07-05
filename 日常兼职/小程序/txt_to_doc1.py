import re

from docx import Document


def txt_to_docx(txt_file, docx_file):
    # 创建一个 Document 对象
    doc = Document()

    # 打开 txt 文件并读取内容
    with open('/Users/hejie/PycharmProjects/turing_exec/日常兼职/小程序/doc.txt', 'r', encoding='utf-8') as file:
        lines = file.readlines()
        pattern = r"(\[.*?题\].*?)(?=正确答案)"
        # 查找所有匹配的文本
        matches = re.findall(pattern, str(lines), re.DOTALL)

        # 输出匹配结果
        for match in matches:
            print(match.strip())
            # doc.add_paragraph(lines.strip())
            # print("\n" + "-" * 40 + "\n")
        # doc.add_paragraph(line.strip())  # 添加段落

    # # 保存为 .docx 文件
    # doc.save(docx_file)
    # print(f"文件已保存为 {docx_file}")


# 示例使用
txt_to_docx('example.txt', 'output.docx')
